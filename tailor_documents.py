#!/usr/bin/env python3
# Tailor Documents - Automated Resume and Cover Letter Generator
# This script processes a template resume and cover letter along with job descriptions
# to create tailored documents for job applications.

import os
import argparse
import json
from datetime import datetime
import re
from pathlib import Path
import shutil

# Import the environment variable loader
try:
    from load_env import main as load_env
    # Try to load environment variables
    load_env()
except ImportError:
    print("Environment loader not found. OpenAI features may not work.")

# We'll use these libraries (will need to be installed via requirements.txt)
try:
    import docx
    from docx import Document
    from PyPDF2 import PdfReader, PdfWriter
except ImportError:
    print("Some required packages are missing. Please run 'pip install -r requirements.txt'")
    exit(1)

# Make OpenAI optional
try:
    import openai
    OPENAI_AVAILABLE = True
except ImportError:
    print("OpenAI package not available. AI features will be disabled.")
    OPENAI_AVAILABLE = False

# Make spaCy optional
try:
    import spacy
    SPACY_AVAILABLE = True
except ImportError:
    print("spaCy not available. Using basic text processing.")
    SPACY_AVAILABLE = False

class DocumentTailorer:
    def __init__(self, config=None):
        self.config = config or {}
        self.input_dir = Path("input")
        self.output_dir = Path("output")
        
        # Initialize NLP model if available
        self.nlp = None
        if SPACY_AVAILABLE:
            try:
                self.nlp = spacy.load("en_core_web_sm")
            except:
                print("spaCy model not found. Using basic text processing.")
            
        # Initialize OpenAI if API key is present
        self.openai_available = False
        if OPENAI_AVAILABLE and "OPENAI_API_KEY" in os.environ:
            openai.api_key = os.environ["OPENAI_API_KEY"]
            self.openai_available = True
        
    def load_resume(self, resume_path):
        """Load a resume file (PDF or DOCX)"""
        resume_path = str(resume_path)  # Convert Path to string
        if resume_path.endswith('.pdf'):
            # Logic for PDF resumes will go here
            return {"type": "pdf", "path": resume_path, "content": "PDF content extraction"}
        elif resume_path.endswith('.docx'):
            doc = Document(resume_path)
            content = "\n".join([para.text for para in doc.paragraphs if para.text])
            return {"type": "docx", "path": resume_path, "content": content, "document": doc}
        elif resume_path.endswith('.txt'):
            with open(resume_path, 'r', encoding='utf-8') as f:
                content = f.read()
            return {"type": "txt", "path": resume_path, "content": content}
        else:
            raise ValueError(f"Unsupported resume format: {resume_path}")
    
    def load_cover_letter(self, cover_letter_path):
        """Load a cover letter template file (PDF or DOCX)"""
        # Similar to load_resume method
        cover_letter_path = str(cover_letter_path)  # Convert Path to string
        if cover_letter_path.endswith('.docx'):
            doc = Document(cover_letter_path)
            content = "\n".join([para.text for para in doc.paragraphs if para.text])
            return {"type": "docx", "path": cover_letter_path, "content": content, "document": doc}
        elif cover_letter_path.endswith('.txt'):
            with open(cover_letter_path, 'r', encoding='utf-8') as f:
                content = f.read()
            return {"type": "txt", "path": cover_letter_path, "content": content}
        else:
            raise ValueError(f"Unsupported cover letter format: {cover_letter_path}")
    
    def load_job_description(self, job_desc_path):
        """Load a job description file"""
        with open(job_desc_path, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Basic info extraction
        job_info = {
            "full_text": content,
            "company": self._extract_company(content),
            "job_title": self._extract_job_title(content),
            "keywords": self._extract_keywords(content)
        }
        
        return job_info
    
    def _extract_company(self, text):
        """Attempt to extract company name from job description"""
        # Simplified approach - could be improved with ML
        company_patterns = [
            r"(?:at|with|for|join)\s+([A-Z][A-Za-z0-9\s&]+)(?:,|\.|is|\n)",
            r"About ([A-Z][A-Za-z0-9\s&]+)(?:,|\.|:|\n)",
            r"([A-Z][A-Za-z0-9\s&]+) is looking for"
        ]
        
        for pattern in company_patterns:
            matches = re.search(pattern, text)
            if matches:
                return matches.group(1).strip()
        
        return "COMPANY_NAME"
    
    def _extract_job_title(self, text):
        """Attempt to extract job title from job description"""
        # Look for common job title patterns
        title_patterns = [
            r"(?:hiring|for|seeking)(?: a| an)? ([A-Za-z]+\s[A-Za-z]+(?:\s[A-Za-z]+)?) (?:to|who|that)",
            r"([A-Za-z]+\s[A-Za-z]+(?:\s[A-Za-z]+)?)(?: position)",
            r"(?:Job Title|Title|Position):?\s*([A-Za-z]+\s[A-Za-z]+(?:\s[A-Za-z]+)?)"
        ]
        
        for pattern in title_patterns:
            matches = re.search(pattern, text)
            if matches:
                return matches.group(1).strip()
        
        return "POSITION_TITLE"
    
    def _extract_keywords(self, text):
        """Extract important keywords from the job description"""
        skills = set()
        
        if self.nlp:
            # Use spaCy if available
            doc = self.nlp(text)
            for ent in doc.ents:
                if ent.label_ in ["ORG", "PRODUCT"]:
                    skills.add(ent.text)
        else:
            # Basic keyword extraction using regex
            # Look for common technical terms
            tech_terms = r'\b(Python|JavaScript|Java|C\+\+|React|Angular|Node\.js|SQL|AWS|Docker|Kubernetes|CI/CD|Machine Learning|Data Science|Project Management)\b'
            matches = re.finditer(tech_terms, text, re.IGNORECASE)
            for match in matches:
                skills.add(match.group(0))
        
        # Add common programming languages, tools, etc.
        tech_keywords = ["Python", "JavaScript", "Java", "C++", "React", "Angular", 
                        "Node.js", "SQL", "AWS", "Docker", "Kubernetes", "CI/CD",
                        "Machine Learning", "Data Science", "Project Management"]
        
        for keyword in tech_keywords:
            if keyword.lower() in text.lower():
                skills.add(keyword)
                
        return list(skills)
    
    def tailor_resume(self, resume_data, job_data):
        """Generate a tailored resume based on the job description"""
        # In a real implementation, this would reorder skills, highlight relevant experience, etc.
        # For now, we'll just make a copy of the original
        
        if self.openai_available:
            prompt = f"""
            Analyze this job description and suggest 5-7 specific modifications to make the resume more relevant:
            
            JOB DESCRIPTION:
            {job_data['full_text'][:2000]}  # Truncated to avoid token limits
            
            RESUME:
            {resume_data['content'][:2000]}  # Truncated to avoid token limits
            
            Please provide specific suggestions in the format:
            1. [SECTION] - [SPECIFIC CHANGE RECOMMENDATION]
            """
            
            try:
                response = openai.ChatCompletion.create(
                    model="gpt-3.5-turbo",
                    messages=[{"role": "user", "content": prompt}]
                )
                suggestions = response.choices[0].message.content
                print("\nTailoring suggestions:\n" + suggestions)
            except Exception as e:
                print(f"OpenAI API call failed: {e}")
                suggestions = "API call failed"
        
        # For a simple MVP, we'll just copy the original file
        if resume_data['type'] == 'docx':
            # Create a new document based on the original
            new_doc = Document()
            for para in resume_data['document'].paragraphs:
                new_para = new_doc.add_paragraph(para.text)
                # Copy formatting (basic)
                for run in para.runs:
                    new_run = new_para.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
            
            return {"type": "docx", "document": new_doc}
        else:
            # For other types, just return the original content for now
            return {"type": "txt", "content": resume_data['content']}
    
    def tailor_cover_letter(self, cover_letter_data, job_data):
        """Generate a tailored cover letter based on the job description"""
        # Replace placeholder text in the cover letter
        content = cover_letter_data['content']
        
        # Replace common placeholders
        replacements = {
            "[COMPANY_NAME]": job_data['company'],
            "[POSITION_TITLE]": job_data['job_title'],
            "[Current Date]": datetime.now().strftime("%B %d, %Y"),
            "[job board/website]": "your job posting",
            "[your field]": "software development",
            "[relevant skills]": ", ".join(job_data['keywords'][:3]),
            "[skill 1]": job_data['keywords'][0] if job_data['keywords'] else "programming",
            "[skill 2]": job_data['keywords'][1] if len(job_data['keywords']) > 1 else "development",
            "[skill 3]": job_data['keywords'][2] if len(job_data['keywords']) > 2 else "problem-solving",
            "[Previous Company]": "my previous company",
            "[brief accomplishment that relates to the job]": f"successfully delivered projects using {job_data['keywords'][0] if job_data['keywords'] else 'various technologies'}",
            "[another accomplishment]": "leading multiple successful project deliveries",
            "[relevant skill for the job]": job_data['keywords'][0] if job_data['keywords'] else "deliver high-quality solutions",
            "[something specific about the company that you admire]": "your focus on innovation and technical excellence",
            "[mention something specific about their products, services, culture, or mission]": "commitment to delivering high-quality software solutions",
            "[relevant experience]": ", ".join(job_data['keywords'][:2]) if job_data['keywords'] else "software development",
            "[specific relevant knowledge]": job_data['keywords'][0] if job_data['keywords'] else "technical expertise",
            "[relevant area]": "software development and innovation",
            "[your phone number]": "YOUR_PHONE_NUMBER",
            "[your email address]": "YOUR_EMAIL_ADDRESS"
        }
        
        for placeholder, value in replacements.items():
            content = content.replace(placeholder, value)
        
        # For a simple MVP, we'll just create a new document with the modified content
        if cover_letter_data['type'] == 'docx':
            # Create a new document based on the original
            new_doc = Document()
            
            # Copy paragraphs from original with replacements
            for para in cover_letter_data['document'].paragraphs:
                para_text = para.text
                for placeholder, value in replacements.items():
                    para_text = para_text.replace(placeholder, value)
                
                new_para = new_doc.add_paragraph(para_text)
                # Copy formatting (basic)
                for run in para.runs:
                    new_run = new_para.add_run(run.text)
                    new_run.bold = run.bold
                    new_run.italic = run.italic
            
            return {"type": "docx", "document": new_doc}
        else:
            # For other types, just return the modified content
            return {"type": "txt", "content": content}

    def save_document(self, doc_data, output_path):
        """Save the tailored document to the output directory"""
        if doc_data['type'] == 'docx':
            doc_data['document'].save(output_path)
            return output_path
        else:
            # For text documents
            with open(output_path, 'w', encoding='utf-8') as f:
                f.write(doc_data['content'])
            return output_path

    def process_job(self, resume_path, cover_letter_path, job_desc_path, job_name):
        """Process a single job application"""
        print(f"Processing job: {job_name}")
        
        # Load input documents
        resume_data = self.load_resume(resume_path)
        cover_letter_data = self.load_cover_letter(cover_letter_path)
        job_data = self.load_job_description(job_desc_path)
        
        # Generate tailored documents
        tailored_resume = self.tailor_resume(resume_data, job_data)
        tailored_cover_letter = self.tailor_cover_letter(cover_letter_data, job_data)
        
        # Create job-specific output directory
        job_output_dir = self.output_dir / job_name
        os.makedirs(job_output_dir, exist_ok=True)
        
        # Save tailored documents - use .txt extension since we're working with text files
        resume_output = job_output_dir / f"Resume_{job_data['company']}_{job_name}.txt"
        cover_letter_output = job_output_dir / f"CoverLetter_{job_data['company']}_{job_name}.txt"
        
        self.save_document(tailored_resume, resume_output)
        self.save_document(tailored_cover_letter, cover_letter_output)
        
        print(f"Tailored documents saved to {job_output_dir}")
        return {
            "resume": str(resume_output),
            "cover_letter": str(cover_letter_output),
            "job_info": job_data
        }
    
    def process_all_jobs(self):
        """Process all job descriptions found in the input directory"""
        # Get base resume and cover letter
        resume_dir = self.input_dir / "resumes"
        cover_letter_dir = self.input_dir / "cover_letters"
        job_desc_dir = self.input_dir / "job_descriptions"
        
        # Find the first resume and cover letter (in a more advanced version, user would select)
        resume_files = list(resume_dir.glob("*.docx")) + list(resume_dir.glob("*.pdf")) + list(resume_dir.glob("*.txt"))
        cover_letter_files = list(cover_letter_dir.glob("*.docx")) + list(cover_letter_dir.glob("*.txt"))
        
        if not resume_files:
            print("No resume files found in input/resumes directory.")
            return
            
        if not cover_letter_files:
            print("No cover letter templates found in input/cover_letters directory.")
            return
            
        # Use the first files found (could be enhanced to let user select)
        resume_path = resume_files[0]
        cover_letter_path = cover_letter_files[0]
        
        print(f"Using resume: {resume_path}")
        print(f"Using cover letter template: {cover_letter_path}")
        
        # Process each job description
        job_desc_files = list(job_desc_dir.glob("*.txt"))
        
        results = []
        for job_file in job_desc_files:
            job_name = job_file.stem
            result = self.process_job(resume_path, cover_letter_path, job_file, job_name)
            results.append(result)
            
        return results

def main():
    parser = argparse.ArgumentParser(description='Tailor resumes and cover letters for job applications')
    parser.add_argument('--config', help='Path to configuration file')
    parser.add_argument('--job', help='Process a specific job (by name or file path)')
    args = parser.parse_args()
    
    config = {}
    if args.config:
        with open(args.config, 'r') as f:
            config = json.load(f)
    
    tailorer = DocumentTailorer(config)
    
    if args.job:
        # Process a single job
        job_path = Path(args.job)
        if not job_path.exists():
            job_path = Path("input/job_descriptions") / f"{args.job}.txt"
        
        if not job_path.exists():
            print(f"Job description not found: {job_path}")
            return
            
        # Use the first resume and cover letter
        resume_dir = Path("input/resumes")
        cover_letter_dir = Path("input/cover_letters")
        
        resume_files = list(resume_dir.glob("*.docx")) + list(resume_dir.glob("*.pdf"))
        cover_letter_files = list(cover_letter_dir.glob("*.docx"))
        
        if not resume_files or not cover_letter_files:
            print("Missing resume or cover letter templates in input directories")
            return
            
        tailorer.process_job(resume_files[0], cover_letter_files[0], job_path, job_path.stem)
    else:
        # Process all jobs
        tailorer.process_all_jobs()

if __name__ == "__main__":
    main() 